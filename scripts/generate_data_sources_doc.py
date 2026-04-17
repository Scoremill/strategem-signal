"""
StrategemSignal Data Sources Reference.

Grid-style Word document cataloging every external data source the v2
app ingests — what it gives us, how often it refreshes, which
pipeline/cron fetches it, which database table(s) it lands in, which
features consume it, known coverage gaps, and source citation.

Primary organization: by data source.
Secondary cross-reference: by feature.

Output lands at project root: StrategemSignal_Data_Sources.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── Paths ────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = ROOT / "StrategemSignal_Data_Sources.docx"

# ── Brand ────────────────────────────────────────────────────────────
ORANGE = RGBColor(0xF9, 0x73, 0x16)
ORANGE_DARK = RGBColor(0xEA, 0x58, 0x0C)
ORANGE_TINT = "FFF7ED"      # shading goes in hex string
DARK_BLUE = RGBColor(0x1E, 0x29, 0x3B)
SLATE_600 = RGBColor(0x4B, 0x55, 0x63)
SLATE_500 = RGBColor(0x6B, 0x72, 0x80)
HEADER_ROW_HEX = "F97316"   # header row fill
ROW_ALT_HEX = "FFFBF5"      # alternating row tint
LIGHT_LINE = "E5E7EB"

FONT = "Calibri"


# ── Shading / cell helpers (python-docx doesn't expose these directly) ──
def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, size=9, bold=False, color=DARK_BLUE,
                  align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    para = cell.paragraphs[0]
    para.alignment = align
    # clear any default run
    for run in list(para.runs):
        run.text = ""
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_cell_borders(cell, color_hex=LIGHT_LINE, sz="4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tc_pr.append(borders)


# ── Document setup ───────────────────────────────────────────────────
doc = Document()
# Landscape is better for wide grids
section = doc.sections[0]
section.page_height, section.page_width = section.page_width, section.page_height
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

# Default font for the whole document
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(11)
style.font.color.rgb = DARK_BLUE


# ── Block-level writers ──────────────────────────────────────────────
def add_title(text: str, subtitle: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = ORANGE
    if subtitle:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(subtitle)
        r2.font.name = FONT
        r2.font.size = Pt(13)
        r2.font.color.rgb = SLATE_600


def add_h1(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = ORANGE


def add_h2(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = ORANGE_DARK


def add_body(text: str, *, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.color.rgb = DARK_BLUE


def add_muted(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(10)
    r.font.color.rgb = SLATE_500


# ── The canonical grid ──────────────────────────────────────────────
# Columns: Source · What It Gives Us · Refresh Cadence · Pipeline / Cron ·
#          Database Table(s) · Consumed By · Coverage Gaps · Source Citation

COLUMN_HEADERS = [
    "Source",
    "What It Gives Us",
    "Refresh Cadence",
    "Pipeline / Cron",
    "Database Table(s)",
    "Consumed By",
    "Coverage Gaps",
    "Source Citation",
]
COLUMN_WIDTHS_CM = [2.4, 3.4, 2.4, 3.0, 2.8, 3.4, 3.0, 4.0]


# ── Helpers for table blocks ─────────────────────────────────────────
def start_grid_table(headers=COLUMN_HEADERS, widths=COLUMN_WIDTHS_CM):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    # Set column widths
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_ROW_HEX)
        set_cell_text(cell, h, size=9, bold=True,
                      color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_borders(cell, "F97316", "6")
    return table


def add_grid_row(table, values, *, row_index=None, shade=False):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        if shade:
            set_cell_shading(cell, ROW_ALT_HEX)
        set_cell_text(cell, v, size=9)
        set_cell_borders(cell)


def add_grid_rows(table, rows):
    """Convenience: add many rows with alternating shading."""
    for i, values in enumerate(rows):
        add_grid_row(table, values, shade=(i % 2 == 1))


# ── DOCUMENT CONTENT ────────────────────────────────────────────────

add_title(
    "StrategemSignal · Data Sources Reference",
    "Every external data feed the v2 app ingests — what, when, where, and who uses it",
)

add_body(
    "This grid-style reference maps every external data source in StrategemSignal v2 "
    "to its refresh cadence, ingestion pipeline, destination table, and downstream "
    "consumers. The goal is one defensible source of truth that answers three "
    "questions at once: where does each number come from, how fresh is it, and which "
    "feature breaks if it goes missing? This document is the companion to the "
    "forthcoming in-app \"View Sources\" modal (Phase 4.5) and the authoritative "
    "reference for anyone onboarding to the codebase.",
)

add_muted(
    "Scope: StrategemSignal v2 only. StrategemOps data is referenced where the v2 "
    "app consumes it via the ops_* mirror tables, but the StrategemOps ingestion "
    "itself is documented separately in that project. Organization is primary by "
    "data source, with a cross-reference table by consuming feature at the end."
)

# ── Section 1 — By Data Source ──────────────────────────────────────
add_h1("1 · By Data Source")

add_body(
    "Each row below is one external data feed. Pipelines are the TypeScript files "
    "in src/lib/pipelines that do the actual fetching; crons are the GitHub Actions "
    "workflow files that fire them. All times are UTC.",
    size=10,
)

# ── 1.1 Federal / Government Sources ────────────────────────────────
add_h2("1.1 · Federal & Government Sources")

federal_rows = [
    # Source, What, Cadence, Pipeline/Cron, Tables, Consumed By, Gaps, Citation
    [
        "Census Bureau · Building Permits Survey (BPS)",
        "Monthly single-family + multi-family residential building permits by metro. Primary signal for home-construction demand. Direct XLS download from Census (not FRED mirror).",
        "Monthly · 6th @ 11:00 UTC",
        "census-bps-client.ts · demand-pipeline.ts · demand-data.yml · /api/cron/demand",
        "permit_data",
        "Portfolio Health Demand sub-score (permits YoY at 30% weight), Filter 3 (Supply-Demand Imbalance), market narratives",
        "197/199 markets covered (2 PR metros excluded). Some metros publish quarterly rather than monthly — permits for those lag by ~30 days.",
        "https://www.census.gov/construction/bps/",
    ],
    [
        "Census Bureau · American Community Survey (ACS) B19013",
        "Median household income by metro. Drives affordability and purchasing-power signals.",
        "Annual · released Dec, pulled annually",
        "income-pipeline.ts · /api/cron/income",
        "income_data",
        "Filter 5 (Affordability Runway), Portfolio Health Financial sub-score (income YoY at 60% weight, level at 40%)",
        "199/199 markets covered. One-year lag is structural; no fix available.",
        "https://www.census.gov/programs-surveys/acs/",
    ],
    [
        "Census Bureau · Population Estimates Program (PEP)",
        "Annual population estimates and net domestic migration by metro. Direct CSV download from Census servers (the API was discontinued for MSAs post-2019, but the same data is published as files).",
        "Annual · new vintage published Dec-March",
        "census-pep-client.ts · demand-pipeline.ts · demand-data.yml",
        "migration_data",
        "Filter 1 (Migration Tailwinds — migration as % of population), Portfolio Health Demand sub-score (population change at 25% weight)",
        "197/199 markets covered (2 PR metros excluded from Census metro CSV). 5 years of annual data per vintage (2020-2024). Cleveland requires CBSA override (17460→17410).",
        "https://www.census.gov/programs-surveys/popest.html",
    ],
    [
        "BLS · Current Employment Statistics (CES)",
        "Total nonfarm employment by metro, month-over-month and year-over-year change. Direct BLS v2 API (not FRED mirror).",
        "Monthly · 6th @ 11:00 UTC",
        "bls-v2-client.ts · demand-pipeline.ts · demand-data.yml · /api/cron/demand",
        "employment_data",
        "Portfolio Health Demand sub-score (employment YoY at 25% weight)",
        "199/199 markets covered. BLS v2 API series ID: SMS{stateFips}{cbsaFips}0000000001. 500 queries/day limit (registered key).",
        "https://www.bls.gov/ces/",
    ],
    [
        "BLS · Local Area Unemployment Statistics (LAUS)",
        "Unemployment rate by metro. Direct BLS v2 API.",
        "Monthly · 6th @ 11:00 UTC",
        "bls-v2-client.ts · demand-pipeline.ts · demand-data.yml",
        "employment_data (unemployment_rate column, same table as CES)",
        "Portfolio Health Demand sub-score (inverted, 20% weight)",
        "198/199 markets covered. BLS v2 API series ID: LAUMT{stateFips}{cbsa4digit}000000003. 1 PR metro lacks LAUS coverage.",
        "https://www.bls.gov/lau/",
    ],
    [
        "BLS · Quarterly Census of Employment & Wages (QCEW)",
        "Construction-trade employment (NAICS 2381/2382/2383/2389) and avg weekly wage by metro. Primary capacity signal. Falls back to state-level data when BLS suppresses MSA-level 238x for smaller metros.",
        "Quarterly · 15th of Jan/Apr/Jul/Oct @ 12:00 UTC",
        "qcew-client.ts · capacity-pipeline.ts · capacity-data.yml · /api/cron/capacity",
        "trade_capacity_data",
        "Business Case Organic model (wage-adjusted build cost), Filter 6 (Operational Feasibility), Portfolio Health Operational sub-score (wage pressure 60%, employment trajectory 40%)",
        "199/199 markets covered (52 metro-level, 147 state-level proxy). State proxy is less precise but honest — same trade wage trends apply statewide. Quarter probe added to avoid targeting unpublished quarters.",
        "https://www.bls.gov/cew/",
    ],
    [
        "BLS · Occupational Employment & Wage Statistics (OEWS, fka OES)",
        "Annual occupation-level employment and wages for construction trades (SOC 47-xxxx).",
        "Annual · April (dual trigger: 13th + 15th)",
        "oes-pipeline.ts · oes-data.yml · /api/cron/oes",
        "occupation_data",
        "Filter 6 (Operational Feasibility), reserved for Phase 4 wage-inflation alerts",
        "Only ~7 of 52 tracked markets covered by OEWS at metro level; QCEW covers all 52 and is preferred",
        "https://www.bls.gov/oes/",
    ],
    [
        "FHFA · House Price Index (HPI)",
        "Quarterly home-price index by metro. Trajectory signal for affordability and market strength.",
        "Monthly · 22nd @ 11:30 UTC (new quarters land in FHFA releases ~60 days after quarter-end)",
        "fhfa-pipeline.ts · fhfa.yml · /api/cron/fhfa",
        "fhfa_hpi",
        "Filter 5 (Affordability Runway) — HPI YoY is the primary trajectory input",
        "No coverage gaps for CBSAs; metro-level FHFA covers 100% of tracked markets",
        "https://www.fhfa.gov/DataTools/Downloads",
    ],
]

table = start_grid_table()
add_grid_rows(table, federal_rows)

# ── 1.2 Commercial Free Sources ─────────────────────────────────────
add_h2("1.2 · Commercial Free Sources")

commercial_rows = [
    [
        "Zillow · ZHVI (Home Value Index)",
        "Monthly median home value in dollars (not an index), metro-level. Drives Business Case sale-price baseline and organic land-cost share.",
        "Monthly · 17th @ 13:00 UTC (Zillow publishes 15th; we give a 48hr buffer)",
        "zillow-zhvi-pipeline.ts · zillow-zhvi.yml · /api/cron/zillow-zhvi",
        "zillow_zhvi",
        "Business Case Organic model (median home price, sale price, land basis), market-tier classifier (A/B/C/D)",
        "195/199 markets covered (98%). 4 unmatched: 2 Puerto Rico metros, Dayton-Kettering OH, Prescott Valley-Prescott AZ.",
        "https://www.zillow.com/research/data/",
    ],
    [
        "FRED (St. Louis Fed)",
        "Legacy fallback for BLS CES/LAUS/population time series. Retained for the original 52 markets with known FRED series IDs. New markets use direct BLS v2 and Census APIs instead.",
        "Legacy — no longer the primary path for any data source",
        "fred-client.ts (library, retained for backward compatibility)",
        "employment_data, migration_data",
        "Indirect — FRED is a transport layer. Direct BLS/Census APIs are now the primary path for all 199 markets.",
        "FRED only covers ~60 top metros. Replaced by direct Census BPS, BLS v2 API, and Census PEP CSV for full coverage.",
        "https://fred.stlouisfed.org/",
    ],
]

table = start_grid_table()
add_grid_rows(table, commercial_rows)

# ── 1.3 Cross-App: StrategemOps Mirror ──────────────────────────────
add_h2("1.3 · Cross-App: StrategemOps Mirror")

add_body(
    "StrategemOps is the sister app that ingests public-homebuilder financial "
    "filings and earnings narratives. StrategemSignal mirrors 14 tables from "
    "StrategemOps monthly through a single bridge (ops-snapshot-pipeline.ts). "
    "No user-facing v2 code ever queries StrategemOps directly — all reads go "
    "through the local ops_* mirror tables. The snapshot has a daily self-heal "
    "cron that only fires if the most recent snapshot failed or is older than "
    "35 days.",
    size=10,
)

ops_rows = [
    [
        "StrategemOps · Companies roster",
        "Public-homebuilder directory: ticker, company name, CIK, exchange, builder category, IR URL.",
        "Monthly · 1st @ 11:00 UTC (+ daily self-heal if stale)",
        "ops-snapshot-pipeline.ts · ops-snapshot.yml · ops-snapshot-retry.yml",
        "ops_companies",
        "Competitive Landscape card (company names beside tickers), Acquisition Targets card",
        "Public builders only (~18 names). Private builders are a Phase 4 pipeline.",
        "StrategemOps app · curly-mud-45701913 Neon project",
    ],
    [
        "StrategemOps · Builder Market Presence",
        "LLM-extracted mapping of which public builders cite which metros in their earnings calls (Filter 4 extractor output). Includes confidence (high/medium/low) and mention count.",
        "Monthly · 1st @ 11:00 UTC",
        "ops-snapshot-pipeline.ts",
        "ops_builder_markets",
        "Filter 4 (Competitive Landscape score), Competitive Landscape UI card, Acquisition Targets UI card",
        "Only covers the ~18 public builders. 326 (builder, market) pairs across 18 builders and 74 markets at last snapshot.",
        "StrategemOps earnings-narrative extractor",
    ],
    [
        "StrategemOps · Financial Periods",
        "Quarterly period metadata (fiscal quarter, calendar quarter, filing date) for all tracked builders.",
        "Monthly · 1st @ 11:00 UTC",
        "ops-snapshot-pipeline.ts",
        "ops_financial_periods",
        "Joined into ops_financial_facts for the Market Health PDF footer citations",
        "Fiscal-year variations between builders are handled but add minor quarter-alignment complexity",
        "StrategemOps EDGAR ingest",
    ],
    [
        "StrategemOps · Financial Facts",
        "Quarterly revenue, gross margin, SG&A, cash flow, backlog, orders, closings by builder and period.",
        "Monthly · 1st @ 11:00 UTC",
        "ops-snapshot-pipeline.ts",
        "ops_financial_facts",
        "Reserved for Phase 3.11 Acquisition rebuild (book value, closings volume per target); peer-benchmark context in future iterations",
        "Subject to EDGAR parsing coverage — Phase 3 gap: some builders disclose lot data only in 10-K, not 8-K",
        "SEC EDGAR via StrategemOps parser",
    ],
    [
        "StrategemOps · Builder Operating KPIs",
        "Per-builder community count, avg sales price, cancellation rate, and other operating metrics.",
        "Monthly · 1st @ 11:00 UTC",
        "ops-snapshot-pipeline.ts",
        "ops_builder_operating_kpis",
        "Reserved for Phase 4 peer-comparison overlays in Portfolio Health",
        "Coverage varies by builder disclosure discipline",
        "SEC EDGAR via StrategemOps parser",
    ],
    [
        "StrategemOps · Incentive Tracking",
        "Promotional incentives (rate buydowns, closing-cost credits) mentioned in earnings calls.",
        "Monthly · 1st @ 11:00 UTC",
        "ops-snapshot-pipeline.ts",
        "ops_incentive_tracking",
        "Reserved for Phase 4 alerts (incentive spike = margin compression signal)",
        "Narrative-derived; numerical precision is approximate",
        "StrategemOps earnings-narrative extractor",
    ],
    [
        "StrategemOps · Management Narratives",
        "Cleaned MD&A and risk-factor text from 10-K / 10-Q filings. Heavy text fields NOT mirrored — we keep only metadata + MD&A + risk factors.",
        "Monthly · 1st @ 11:00 UTC",
        "ops-snapshot-pipeline.ts",
        "ops_management_narratives",
        "Reserved for Phase 4 natural-language market-commentary features",
        "Parsed from semi-structured filing HTML; occasional formatting artifacts",
        "SEC EDGAR via StrategemOps parser",
    ],
    [
        "StrategemOps · Innovation Theme Mentions",
        "LLM-tagged mentions of innovation themes (BTR, modular, energy efficiency, AI, etc.) in earnings calls.",
        "Monthly · 1st @ 11:00 UTC",
        "ops-snapshot-pipeline.ts",
        "ops_innovation_theme_mentions",
        "Reserved for future strategic-scanning features",
        "Theme taxonomy is curated manually; gaps reflect that",
        "StrategemOps earnings-narrative extractor",
    ],
    [
        "StrategemOps · Sentiment Scores",
        "Per-builder, per-quarter sentiment scores from earnings-call analysis.",
        "Monthly · 1st @ 11:00 UTC",
        "ops-snapshot-pipeline.ts",
        "ops_sentiment_scores",
        "Reserved — StrategemSignal does not currently surface sentiment. Kept in mirror for Phase 4 features.",
        "Currently scored with gpt-4.1-mini; upgrade to gpt-4.1 is a tracked deferral per memory",
        "StrategemOps sentiment pipeline",
    ],
    [
        "StrategemOps · Sector Sentiment Composite",
        "Rolled-up sector-level sentiment (aggregated across builders).",
        "Monthly · 1st @ 11:00 UTC",
        "ops-snapshot-pipeline.ts",
        "ops_sector_sentiment_composite",
        "Reserved for Phase 4",
        "Same as sentiment_scores",
        "StrategemOps sentiment pipeline",
    ],
    [
        "StrategemOps · Earnings Calendar",
        "Upcoming earnings-call dates for tracked builders.",
        "Monthly · 1st @ 11:00 UTC",
        "ops-snapshot-pipeline.ts",
        "ops_earnings_calendar",
        "Reserved for Phase 4 alerts",
        "Self-populating; occasional late-date corrections by builders",
        "StrategemOps earnings-calendar pipeline",
    ],
    [
        "StrategemOps · Benchmark Ranges",
        "Industry benchmark bands (low / mid / high) for operating and financial metrics.",
        "Full-replace on snapshot (small table)",
        "ops-snapshot-pipeline.ts",
        "ops_benchmark_ranges",
        "Reserved for Phase 4 Portfolio Health band overlays",
        "Curated manually in StrategemOps admin",
        "StrategemOps benchmark admin",
    ],
    [
        "StrategemOps · Source Documents",
        "Metadata for source filings (NOT the text body — kept in StrategemOps).",
        "Full-replace on snapshot",
        "ops-snapshot-pipeline.ts",
        "ops_source_documents",
        "View Sources modal (Phase 4.5) for traceability",
        "Text body intentionally NOT mirrored to keep StrategemSignal DB small",
        "SEC EDGAR",
    ],
    [
        "StrategemOps · Filings metadata",
        "Filing-level metadata (ticker, form type, accession number, filed date) — text body NOT mirrored.",
        "Full-replace on snapshot",
        "ops-snapshot-pipeline.ts",
        "ops_filings",
        "View Sources modal (Phase 4.5)",
        "Text body in StrategemOps only",
        "SEC EDGAR",
    ],
    [
        "StrategemOps · Company Universe Registry",
        "Active/inactive status of tracked companies; controls which appear in the roster.",
        "Monthly · 1st @ 11:00 UTC",
        "ops-snapshot-pipeline.ts",
        "ops_company_universe_registry",
        "Gates which companies show up in ops_companies for v2 queries",
        "None",
        "StrategemOps roster admin",
    ],
]

table = start_grid_table()
add_grid_rows(table, ops_rows)

# ── 1.4 AI / LLM services ──────────────────────────────────────────
add_h2("1.4 · AI & LLM Services")

ai_rows = [
    [
        "OpenAI · gpt-4.1",
        "Generates the per-market narrative blurbs (Portfolio Health blurb + Market Opportunity blurb). Structured JSON output with explicit guardrails to never cite a composite score.",
        "Monthly · 23rd @ 12:00 UTC (after scoring runs)",
        "market-narratives-pipeline.ts · market-narratives.yml · /api/cron/market-narratives",
        "market_narratives",
        "Market drilldown page narrative blocks, heatmap hover popup, PDF export Market Health section",
        "Generates ~199 narratives per monthly run; failures retry within run but do not block the scoring pipeline",
        "OpenAI API; prompt-locked to preset-neutral language per Phase 2 decision",
    ],
]

table = start_grid_table()
add_grid_rows(table, ai_rows)

# ── 1.5 Computed Snapshot Tables ────────────────────────────────────
add_h2("1.5 · Computed Snapshot Tables (not external data, but primary read surface)")

add_body(
    "These tables are not ingested from external sources — they are computed "
    "from the upstream federal + StrategemOps data above. Listed here because "
    "the user-facing app reads from them directly, so it's where most "
    "\"View Sources\" trails will terminate for UI purposes.",
    size=10,
)

computed_rows = [
    [
        "Portfolio Health Snapshot",
        "Monthly per-market Financial / Demand / Operational sub-scores (0-100) + JSON inputs tuple for every data point that fed them.",
        "Monthly · 22nd @ 11:00 UTC",
        "portfolio-health-pipeline.ts · portfolio-health.yml · /api/cron/portfolio-health",
        "portfolio_health_snapshots",
        "Heatmap, Markets list, Market drilldown page, Business Case PDF Market Health section",
        "Derived from upstream data — gaps in upstream feeds (e.g. permit lag) propagate as null sub-scores",
        "Internal — scorer in src/lib/scoring/portfolio-health.ts",
    ],
    [
        "Market Opportunity Scores",
        "Monthly per-market six-filter scores (Migration + Demand, Employment Diversity, Supply-Demand Imbalance, Competitive Landscape, Affordability Runway, Operational Feasibility) + JSON input traces.",
        "Monthly · 22nd @ 11:15 UTC",
        "market-opportunity-pipeline.ts · market-opportunity.yml · /api/cron/market-opportunity",
        "market_opportunity_scores",
        "Markets list + drilldown per-filter views; Phase 4 ranking views",
        "Phase 2 Filter 5 was FHFA-stubbed until FHFA pipeline landed; now live",
        "Internal — scorer in src/lib/scoring/market-opportunity.ts",
    ],
    [
        "Fetch Logs",
        "One row per pipeline run: status, duration, row counts, error messages. Used by the daily self-heal and the admin health page.",
        "Written by every pipeline on every run",
        "All pipelines",
        "fetch_logs",
        "Admin health page, ops-snapshot-retry cron",
        "N/A — diagnostic table only",
        "Internal",
    ],
]

table = start_grid_table()
add_grid_rows(table, computed_rows)

# ── 1.6 Tenant / User Data ──────────────────────────────────────────
add_h2("1.6 · Tenant & User Data (not external; for completeness)")

add_muted(
    "These tables are user-generated, not ingested. Listed for completeness "
    "since they appear in the View Sources modal trail."
)

tenant_rows = [
    [
        "Orgs, Users, Memberships",
        "Multi-tenant foundation: which user belongs to which org, with what role.",
        "On signup / settings change",
        "Auth flow, /api/auth/login, settings actions",
        "orgs, users, org_memberships",
        "Auth, tenant isolation, admin UI",
        "N/A",
        "Internal — created via signup + admin invites",
    ],
    [
        "Tracked + Watchlist Markets",
        "Per-user market filter (which metros this user cares about) and watchlist (flagged for alerts).",
        "On user action",
        "Settings actions · /settings",
        "tracked_markets, watchlist_markets",
        "Markets list filter, heatmap scope, Phase 4 alerts",
        "N/A",
        "Internal",
    ],
    [
        "Health Score Weights",
        "Per-user weighting preset for blending Financial / Demand / Operational into composite. One row per (user, org).",
        "On user action",
        "Settings actions",
        "health_score_weights",
        "Heatmap composite coloring, markets table composite column",
        "N/A — client recomputes composite from stored sub-scores on every view",
        "Internal",
    ],
    [
        "Flags, Business Cases, Alert Preferences, Alerts, Audit Log",
        "Per-user workflow artifacts.",
        "On user action",
        "Various server actions",
        "flags, business_cases, alert_preferences, alerts, audit_log",
        "Market drilldown flags, business case library, Phase 4 alerts, compliance audit",
        "N/A",
        "Internal",
    ],
]

table = start_grid_table()
add_grid_rows(table, tenant_rows)

# ── Section 2 — Cross-Reference by Feature ──────────────────────────
doc.add_page_break()
add_h1("2 · Cross-Reference by Feature")

add_body(
    "The inverse view of Section 1: for each major user-facing feature, which "
    "data sources feed it. Useful for impact analysis when a source goes down "
    "or changes format.",
    size=10,
)

feature_rows = [
    [
        "Heatmap (Portfolio Health map)",
        "Monthly refresh",
        "Census BPS, Census ACS, Census PEP, BLS CES, BLS LAUS, BLS QCEW, OpenAI narratives",
        "permit_data, income_data, migration_data, employment_data, trade_capacity_data, portfolio_health_snapshots, market_narratives",
    ],
    [
        "Markets list + drilldown",
        "Monthly refresh",
        "All Portfolio Health + all Market Opportunity sources (same as the scoring pipelines)",
        "portfolio_health_snapshots, market_opportunity_scores, market_narratives",
    ],
    [
        "Business Case · Organic Entry model",
        "Live re-run per slider move (in-browser, pure function)",
        "Zillow ZHVI (median home price), BLS QCEW (construction wage → regional build-cost multiplier)",
        "zillow_zhvi, trade_capacity_data, geographies",
    ],
    [
        "Business Case · Competitive Landscape card",
        "Monthly via ops-snapshot",
        "StrategemOps Builder Market Presence (from earnings-narrative extractor)",
        "ops_builder_markets, ops_companies",
    ],
    [
        "Business Case · Acquisition Targets card",
        "Monthly via ops-snapshot",
        "Same as Competitive Landscape",
        "ops_builder_markets, ops_companies",
    ],
    [
        "Business Case · Market Health section (PDF)",
        "Monthly refresh",
        "Everything that feeds Portfolio Health Financial/Demand/Operational sub-scores",
        "portfolio_health_snapshots, market_narratives",
    ],
    [
        "Filter 1 · Migration + Demand",
        "Monthly",
        "Census BPS, Census PEP, BLS CES, BLS LAUS",
        "permit_data, migration_data, employment_data",
    ],
    [
        "Filter 2 · Employment Diversity",
        "Quarterly (QCEW cadence)",
        "BLS QCEW 2-digit NAICS sector breakdown (fetched live in the scoring pipeline)",
        "trade_capacity_data (base), plus live BLS QCEW sector call",
    ],
    [
        "Filter 3 · Supply-Demand Imbalance",
        "Monthly",
        "Census BPS, Census PEP, derived permits-per-capita",
        "permit_data, migration_data",
    ],
    [
        "Filter 4 · Competitive Landscape",
        "Monthly (ops-snapshot)",
        "StrategemOps Builder Market Presence",
        "ops_builder_markets",
    ],
    [
        "Filter 5 · Affordability Runway",
        "Monthly",
        "Census ACS (income), FHFA HPI (home-price trajectory)",
        "income_data, fhfa_hpi",
    ],
    [
        "Filter 6 · Operational Feasibility",
        "Quarterly (QCEW cadence) + annual (OEWS)",
        "BLS QCEW (primary), BLS OEWS (supplemental where coverage exists)",
        "trade_capacity_data, occupation_data",
    ],
    [
        "Alerts (Phase 4)",
        "To be defined",
        "All upstream — decision-framed alerts fire on state changes across any tracked signal",
        "alerts + any signal table",
    ],
    [
        "View Sources modal (Phase 4.5)",
        "Read-through on demand",
        "All tables above; traceability is the feature",
        "fetch_logs + originating tables",
    ],
]

# Different column set for cross-reference — narrower layout
CROSS_HEADERS = ["Feature", "Refresh Cadence", "Data Sources", "Tables Read"]
CROSS_WIDTHS = [5.5, 3.5, 8.0, 7.5]

table = start_grid_table(CROSS_HEADERS, CROSS_WIDTHS)
add_grid_rows(table, feature_rows)

# ── Section 3 — Coverage gaps summary ───────────────────────────────
add_h1("3 · Known Coverage Gaps")

add_body(
    "Rolled up from the Coverage Gaps column of every source table. This is "
    "the running list of places we know the data is incomplete — useful for "
    "prioritizing Phase 4 pipeline work and for setting honest expectations "
    "when defending a number.",
    size=10,
)

gap_rows = [
    [
        "Puerto Rico metros (2 markets)",
        "Aguadilla-Isabela PR and San Juan-Bayamón PR lack Census PEP population data, Zillow ZHVI, and FHFA HPI. Census excludes PR from the metro-level PEP CSV.",
        "Scores computed from available inputs only (employment, income, QCEW). Missing components are null, not imputed.",
        "Accept as structural — these markets have limited federal data coverage by design.",
    ],
    [
        "Zillow ZHVI · 4 unmatched markets",
        "2 PR metros, Dayton-Kettering OH, Prescott Valley-Prescott AZ. No Zillow coverage.",
        "Business Case Organic model cannot compute for these markets (returns empty result with a warning).",
        "Accept for PR. Dayton/Prescott may resolve with CBSA code matching update.",
    ],
    [
        "QCEW trade data · state-level proxy for 147 markets",
        "BLS suppresses NAICS 238x (4-digit construction trades) for smaller metros due to employer disclosure rules. 147 of 199 markets use state-level trade wages as proxy.",
        "State-level proxy is clearly sourced ('bls_qcew' in the data). Same trade wage trends apply statewide — less precise than metro-level but better than null.",
        "Accept as structural. Metro-level 238x is only published for ~50 largest MSAs.",
    ],
    [
        "Private builder market presence",
        "Competitive Landscape card covers only ~18 public builders via earnings narratives. Regional + private builders invisible.",
        "Coverage footer on the card flags the gap. Card is public-only by design.",
        "Phase 5.4 — private-builder pipeline (data source decision pending).",
    ],
    [
        "Acquisition cost modeling",
        "Acquisition Targets card is a plain list — no per-builder cost math.",
        "Clearly labeled. Prior 'goodwill premium per unit' was retired in Phase 3.10b as misleading.",
        "Phase 5.5 — total-cost-of-entry model. Depends on 5.4.",
    ],
    [
        "OEWS coverage limits",
        "Only ~7 of 199 tracked markets have OEWS data at the metro level.",
        "Filter 6 uses QCEW as primary; OEWS is supplemental where available.",
        "Accept as structural — BLS does not publish metro-level OEWS for smaller metros.",
    ],
    [
        "ACS income lag",
        "One-year structural lag on median household income.",
        "Documented; Filter 5 operates on the latest available vintage.",
        "No fix available. Accept.",
    ],
    [
        "Census PEP population lag",
        "~18 month lag. Latest available is July 2024 estimates (released Dec 2024).",
        "Annual frequency is appropriate for population — it doesn't change fast enough to need monthly updates.",
        "Accept as structural. Census PEP is the authoritative source.",
    ],
    [
        "EDGAR parsing — lot data",
        "Some builders disclose lot positions only in 10-K (annual), not 8-K (quarterly).",
        "Competitive Landscape data is narrative-derived and unaffected.",
        "StrategemOps parser enhancement (cross-project).",
    ],
    [
        "Sentiment scoring model",
        "ops_sentiment_scores still on gpt-4.1-mini; upgrade to gpt-4.1 is deferred.",
        "Not currently surfaced in StrategemSignal.",
        "Cross-project — StrategemOps retro-score initiative.",
    ],
]

GAP_HEADERS = ["Gap", "Description", "Current Handling", "Resolution Path"]
GAP_WIDTHS = [4.5, 6.0, 5.5, 8.5]

table = start_grid_table(GAP_HEADERS, GAP_WIDTHS)
add_grid_rows(table, gap_rows)

# ── Footer note ─────────────────────────────────────────────────────
doc.add_paragraph()
add_muted(
    "This document is maintained alongside the code. When a new pipeline lands "
    "or an existing source changes cadence, update the corresponding row and "
    "re-run scripts/generate_data_sources_doc.py. The canonical copy lives at "
    "StrategemSignal_Data_Sources.docx in the project root."
)

doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
